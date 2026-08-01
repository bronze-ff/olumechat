from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "output" / "documents" / "Proposta_Comercial_Olume_Chat_Pro_Distribuidora.docx"
LOGO = ROOT / "docs" / "brand" / "assets" / "olume-chat-logo.png"

INK = "071A15"
INK_MUTED = "4D625C"
PRIMARY = "1F7A60"
PRIMARY_DARK = "17664F"
PRIMARY_SOFT = "EAFBF5"
SIGNAL = "5BD6AE"
CANVAS = "F3F8F6"
SURFACE_SUBTLE = "E9EFEA"
BORDER = "D3E0DB"
BORDER_STRONG = "BFD0CA"
WHITE = "FFFFFF"
WARNING_SOFT = "FFF7E8"
WARNING = "B76A11"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 100, "bottom": 100, "start": 140, "end": 140}


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_run_font(run, size=None, color=INK, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, **kwargs):
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    tcmar = tcpr.first_child_found_in("w:tcMar")
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for margin in ["top", "start", "bottom", "end"]:
        node = tcmar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tcmar.append(node)
        node.set(qn("w:w"), str(kwargs.get(margin, CELL_MARGINS_DXA[margin])))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=BORDER, size="6", **edges):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = tcpr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcpr.append(borders)
    for edge in ["top", "start", "bottom", "end", "insideH", "insideV"]:
        edge_data = edges.get(edge, {"val": "single", "color": color, "sz": size})
        if edge_data is None:
            continue
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), edge_data.get("val", "single"))
        tag.set(qn("w:sz"), edge_data.get("sz", size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), edge_data.get("color", color))


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    total = sum(widths_dxa)
    tbl = table._tbl
    tblpr = tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(total))
    tblw.set(qn("w:type"), "dxa")

    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(indent_dxa))
    tblind.set(qn("w:type"), "dxa")

    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)


def mark_header_row(row):
    trpr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trpr.append(header)


def set_repeat_keep(paragraph, keep_with_next=False, keep_together=True):
    paragraph.paragraph_format.keep_together = keep_together
    paragraph.paragraph_format.keep_with_next = keep_with_next


def add_para(doc, text="", size=11, color=INK, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=8,
             line_spacing=1.333, style=None, keep_with_next=False):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line_spacing
    set_repeat_keep(p, keep_with_next=keep_with_next)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_mixed_para(doc, segments, before=0, after=8, line_spacing=1.333,
                   align=WD_ALIGN_PARAGRAPH.LEFT, shading=None, border_color=None,
                   left_indent=0, right_indent=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.left_indent = Inches(left_indent)
    p.paragraph_format.right_indent = Inches(right_indent)
    set_repeat_keep(p)
    for text, kwargs in segments:
        run = p.add_run(text)
        set_run_font(run, **kwargs)
    ppr = p._p.get_or_add_pPr()
    if shading:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), shading)
        ppr.append(shd)
    if border_color:
        pbdr = OxmlElement("w:pBdr")
        for edge in ["top", "bottom", "start", "end"]:
            node = OxmlElement(f"w:{edge}")
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), "6")
            node.set(qn("w:space"), "6")
            node.set(qn("w:color"), border_color)
            pbdr.append(node)
        ppr.append(pbdr)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    set_repeat_keep(p, keep_with_next=True)
    return p


def add_bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    set_repeat_keep(p)
    if bold_lead and text.startswith(bold_lead):
        first = p.add_run(bold_lead)
        set_run_font(first, size=11, color=INK, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, size=11, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11, color=INK)
    return p


def add_numbered(doc, title, body):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.208
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    set_repeat_keep(p)
    r1 = p.add_run(f"{title}. ")
    set_run_font(r1, size=11, color=INK, bold=True)
    r2 = p.add_run(body)
    set_run_font(r2, size=11, color=INK)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_text, fld_end])
    set_run_font(run, size=9, color=INK_MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = rgb(INK)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = rgb(INK_MUTED)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(16)

    heading_tokens = {
        1: (16, PRIMARY, 18, 10),
        2: (13, PRIMARY, 12, 6),
        3: (12, PRIMARY_DARK, 8, 4),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_section(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("PROPOSTA OLU/PRO/2026-0731")
    set_run_font(r, size=8.5, color=INK_MUTED, bold=True)

    footer = section.footer
    ft = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    ft.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(ft, [7200, 2160], indent_dxa=0)
    for c in ft.rows[0].cells:
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(c, top=20, bottom=20, start=0, end=0)
        set_cell_border(c, edges={})
    p0 = ft.cell(0, 0).paragraphs[0]
    p0.paragraph_format.space_after = Pt(0)
    r0 = p0.add_run("olumechat.com.br  |  comercial@olumechat.com.br")
    set_run_font(r0, size=8.5, color=INK_MUTED)
    p1 = ft.cell(0, 1).paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p1.paragraph_format.space_after = Pt(0)
    rr = p1.add_run("Página ")
    set_run_font(rr, size=8.5, color=INK_MUTED)
    add_page_field(p1)


def set_cell_text(cell, text, size=10.5, color=INK, bold=False,
                  align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.keep_together = True
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_grid_table(table, header=True, header_fill=PRIMARY_SOFT):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_margins(cell)
            set_cell_border(cell)
            if ri == 0 and header:
                set_cell_shading(cell, header_fill)
        if ri == 0 and header:
            mark_header_row(row)


def build_document():
    doc = Document()
    configure_styles(doc)
    configure_section(doc)

    props = doc.core_properties
    props.title = "Proposta Comercial Olume Chat para Pro Distribuidora"
    props.subject = "Implantação da plataforma Olume Chat e integrações opcionais com Winthor"
    props.author = "Olume"
    props.keywords = "Olume Chat, WhatsApp, Meta, Winthor, proposta comercial"
    props.comments = "Documento comercial emitido em 31/07/2026."

    # PAGE 1 - PROPOSAL CENTERPIECE
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_after = Pt(20)
    shape = p_logo.add_run().add_picture(str(LOGO), width=Inches(2.15))
    shape._inline.docPr.set("descr", "Logotipo Olume Chat")

    add_para(doc, "PROPOSTA COMERCIAL", size=10, color=PRIMARY, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=7, line_spacing=1.0)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Atendimento no WhatsApp com operação centralizada e custo previsível")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Proposta preparada para a Pro Distribuidora")

    meta = doc.add_table(rows=2, cols=2)
    set_table_geometry(meta, [4680, 4680])
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Aos cuidados de", "Iure"),
        ("Data", "31 de julho de 2026"),
        ("Proposta", "OLU/PRO/2026-0731"),
        ("Validade", "15 dias"),
    ]
    for idx, (label, value) in enumerate(meta_data):
        cell = meta.cell(idx // 2, idx % 2)
        set_cell_shading(cell, CANVAS)
        set_cell_border(cell, color=BORDER)
        set_cell_margins(cell, top=120, bottom=120, start=150, end=150)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(label.upper())
        set_run_font(r1, size=8, color=INK_MUTED, bold=True)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(value)
        set_run_font(r2, size=10.5, color=INK, bold=True)

    add_para(doc, "Uma proposta simples de contratar e simples de acompanhar", size=15,
             color=INK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=22, after=6)
    add_para(
        doc,
        "A Olume Chat reúne atendimento, equipe e automações no WhatsApp oficial. "
        "A Pro Distribuidora paga uma única mensalidade pela plataforma; integrações com o "
        "Winthor entram somente na implantação, de acordo com a dificuldade de cada item.",
        size=11, color=INK_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=18,
    )

    summary = doc.add_table(rows=2, cols=3)
    set_table_geometry(summary, [3120, 3120, 3120])
    labels = ["IMPLANTAÇÃO BASE", "MENSALIDADE", "API OFICIAL DA META"]
    values = ["R$ 500", "R$ 500/mês", "Conforme consumo"]
    for col in range(3):
        top = summary.cell(0, col)
        bot = summary.cell(1, col)
        set_cell_text(top, labels[col], size=8.2, color=PRIMARY, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(bot, values[col], size=15 if col < 2 else 12.5, color=INK,
                      bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(top, PRIMARY_SOFT)
        set_cell_shading(bot, PRIMARY_SOFT)
        set_cell_border(top, color=BORDER_STRONG)
        set_cell_border(bot, color=BORDER_STRONG)
        set_cell_margins(top, top=100, bottom=60, start=90, end=90)
        set_cell_margins(bot, top=50, bottom=120, start=90, end=90)

    add_mixed_para(
        doc,
        [
            ("Ponto central: ", {"size": 10.5, "color": PRIMARY_DARK, "bold": True}),
            ("as integrações com o Winthor não geram mensalidade adicional. ", {"size": 10.5, "color": INK}),
            ("O recorrente permanece R$ 500 + custo da API oficial da Meta.", {"size": 10.5, "color": INK, "bold": True}),
        ],
        before=16,
        after=0,
        shading=PRIMARY_SOFT,
        border_color=BORDER_STRONG,
        left_indent=0.08,
        right_indent=0.08,
    )

    add_page_break(doc)

    # PAGE 2 - SCOPE AND IMPLEMENTATION
    add_heading(doc, "1. O que a Pro Distribuidora recebe", 1)
    add_para(
        doc,
        "Uma operação de WhatsApp baseada na Cloud API oficial da Meta, pronta para organizar "
        "o atendimento da equipe e automatizar etapas repetitivas sem perder o contexto da conversa.",
        size=11, color=INK_MUTED, after=10,
    )
    add_bullet(doc, "Caixa de entrada compartilhada para centralizar conversas e histórico.", "Caixa de entrada compartilhada")
    add_bullet(doc, "Distribuição do atendimento por filas, departamentos e responsáveis.", "Distribuição do atendimento")
    add_bullet(doc, "Fluxos automáticos de recepção, triagem, menus e regras de encaminhamento.", "Fluxos automáticos")
    add_bullet(doc, "Campanhas e comunicação ativa conforme as políticas da Meta e o consentimento dos contatos.", "Campanhas e comunicação ativa")
    add_bullet(doc, "Gestão de contatos, tags, acompanhamento operacional e visão de métricas.", "Gestão de contatos")
    add_bullet(doc, "Suporte à configuração e às rotinas da plataforma durante a vigência.", "Suporte à configuração")

    add_heading(doc, "2. Implantação base", 1)
    add_para(
        doc,
        "A implantação base custa R$ 500 e prepara a operação para o primeiro uso. O trabalho "
        "começa após o aceite e a liberação dos acessos necessários.",
        size=11, after=10,
    )
    add_numbered(doc, "Kickoff", "alinhamento do responsável, número de WhatsApp e prioridades da operação")
    add_numbered(doc, "Conexão oficial", "orientação e configuração dos ativos necessários na Meta")
    add_numbered(doc, "Configuração", "estrutura inicial de equipe, filas e fluxo de entrada")
    add_numbered(doc, "Go-live", "validação com o cliente e início do acompanhamento assistido")

    add_mixed_para(
        doc,
        [
            ("Prazo estimado: ", {"size": 11, "color": PRIMARY_DARK, "bold": True}),
            ("até 30 dias a partir do aceite, da aprovação do escopo e da entrega dos acessos. ", {"size": 11, "color": INK}),
            ("Após o go-live, a Olume acompanha os ajustes iniciais por 30 dias.", {"size": 11, "color": INK, "bold": True}),
        ],
        before=10,
        after=14,
        shading=CANVAS,
        border_color=BORDER,
        left_indent=0.08,
        right_indent=0.08,
    )

    add_heading(doc, "3. Responsabilidades para começar", 1)
    add_bullet(doc, "A Pro Distribuidora indica um responsável para decisões, validações e treinamento.")
    add_bullet(doc, "O cliente disponibiliza número, conta Meta e acessos necessários ao ambiente Winthor quando houver integração.")
    add_bullet(doc, "Cada item de integração é descrito, classificado e aprovado antes do desenvolvimento.")
    add_bullet(doc, "Dependências ou aprovações externas da Meta e do Winthor podem alterar o cronograma.")

    add_page_break(doc)

    # PAGE 3 - PRICING
    add_heading(doc, "4. Investimento", 1)
    add_para(
        doc,
        "A cobrança foi estruturada para separar claramente o custo inicial do custo mensal. "
        "A integração é investimento de implantação; ela não vira recorrência.",
        size=11, color=INK_MUTED, after=10,
    )

    base = doc.add_table(rows=4, cols=3)
    set_table_geometry(base, [4200, 2280, 2880])
    headers = ["Item", "Cobrança", "Valor"]
    for i, h in enumerate(headers):
        set_cell_text(base.cell(0, i), h, size=9.5, color=PRIMARY_DARK, bold=True)
    rows = [
        ("Implantação da plataforma", "Pagamento único", "R$ 500"),
        ("Assinatura Olume Chat", "Mensal", "R$ 500/mês"),
        ("API oficial do WhatsApp", "Conforme consumo", "Tabela vigente da Meta"),
    ]
    for ri, row in enumerate(rows, start=1):
        for ci, value in enumerate(row):
            set_cell_text(base.cell(ri, ci), value, size=10.5,
                          bold=(ci == 2 and ri < 3),
                          align=WD_ALIGN_PARAGRAPH.RIGHT if ci == 2 else WD_ALIGN_PARAGRAPH.LEFT)
    style_grid_table(base)

    add_heading(doc, "Integrações opcionais com o Winthor", 2)
    add_para(
        doc,
        "Cada automação ou troca de dados é tratada como um item de integração. O valor é "
        "somado à implantação conforme o nível de dificuldade validado com o cliente.",
        size=10.8, after=8,
    )
    wint = doc.add_table(rows=4, cols=3)
    set_table_geometry(wint, [1800, 5760, 1800])
    for i, h in enumerate(["Nível", "Referência para classificação", "Valor por item"]):
        set_cell_text(wint.cell(0, i), h, size=9.2, color=PRIMARY_DARK, bold=True,
                      align=WD_ALIGN_PARAGRAPH.RIGHT if i == 2 else WD_ALIGN_PARAGRAPH.LEFT)
    levels = [
        ("Baixa", "Consulta simples, uma fonte principal e resposta direta", "R$ 300"),
        ("Média", "Mais de uma validação, regra condicional ou combinação de dados", "R$ 600"),
        ("Alta", "Fluxo em várias etapas, gravação de dados ou regra crítica de negócio", "R$ 1.000"),
    ]
    for ri, row in enumerate(levels, start=1):
        for ci, value in enumerate(row):
            set_cell_text(wint.cell(ri, ci), value, size=10.2,
                          color=PRIMARY_DARK if ci == 0 else INK,
                          bold=(ci in [0, 2]),
                          align=WD_ALIGN_PARAGRAPH.RIGHT if ci == 2 else WD_ALIGN_PARAGRAPH.LEFT)
    style_grid_table(wint)

    add_heading(doc, "Exemplos de composição", 2)
    examples = doc.add_table(rows=4, cols=3)
    set_table_geometry(examples, [5100, 2460, 1800])
    for i, h in enumerate(["Cenário", "Cálculo", "Implantação total"]):
        set_cell_text(examples.cell(0, i), h, size=9.2, color=PRIMARY_DARK, bold=True,
                      align=WD_ALIGN_PARAGRAPH.RIGHT if i == 2 else WD_ALIGN_PARAGRAPH.LEFT)
    example_rows = [
        ("Sem integração com Winthor", "R$ 500", "R$ 500"),
        ("3 integrações de baixa dificuldade", "R$ 500 + 3 x R$ 300", "R$ 1.400"),
        ("1 baixa + 1 alta", "R$ 500 + R$ 300 + R$ 1.000", "R$ 1.800"),
    ]
    for ri, row in enumerate(example_rows, start=1):
        for ci, value in enumerate(row):
            set_cell_text(examples.cell(ri, ci), value, size=10.2,
                          bold=(ci == 2),
                          align=WD_ALIGN_PARAGRAPH.RIGHT if ci == 2 else WD_ALIGN_PARAGRAPH.LEFT)
            if ri == 2:
                set_cell_shading(examples.cell(ri, ci), PRIMARY_SOFT)
    style_grid_table(examples)

    add_mixed_para(
        doc,
        [
            ("Outros sistemas: ", {"size": 10.8, "color": WARNING, "bold": True}),
            ("integrações fora do Winthor passam por validação técnica e recebem orçamento específico antes do início.", {"size": 10.8, "color": INK}),
        ],
        before=12,
        after=0,
        shading=WARNING_SOFT,
        border_color="E8D4AA",
        left_indent=0.08,
        right_indent=0.08,
    )

    add_page_break(doc)

    # PAGE 4 - CONDITIONS AND ACCEPTANCE
    add_heading(doc, "5. Condições comerciais", 1)
    add_bullet(doc, "Valores apresentados em reais (R$).")
    add_bullet(doc, "Implantação e integrações: cobrança única após o aceite comercial e a aprovação do escopo técnico.")
    add_bullet(doc, "Mensalidade: R$ 500, com início no go-live da plataforma.")
    add_bullet(doc, "Consumo da API oficial: não incluído na mensalidade; calculado conforme a política e a tabela vigentes da Meta.")
    add_bullet(doc, "Integrações Winthor: cobradas por item e apenas na implantação, sem recorrência adicional.")
    add_bullet(doc, "Mudanças de escopo e novas integrações após a aprovação recebem estimativa separada.")
    add_bullet(doc, "Validade desta proposta: 15 dias a partir de 31 de julho de 2026.")

    add_heading(doc, "6. Próximos passos", 1)
    add_numbered(doc, "Aceite comercial", "confirmação desta proposta pela Pro Distribuidora")
    add_numbered(doc, "Levantamento técnico", "lista dos itens Winthor e classificação de dificuldade")
    add_numbered(doc, "Ordem de serviço", "registro do escopo final, valores e dados de faturamento")
    add_numbered(doc, "Kickoff", "agendamento do início e compartilhamento dos acessos")

    add_mixed_para(
        doc,
        [
            ("Resumo para decisão: ", {"size": 11, "color": PRIMARY_DARK, "bold": True}),
            ("R$ 500 de implantação base + integrações escolhidas. Depois do go-live, ", {"size": 11, "color": INK}),
            ("R$ 500 por mês + custo da API oficial da Meta.", {"size": 11, "color": INK, "bold": True}),
        ],
        before=12,
        after=18,
        shading=PRIMARY_SOFT,
        border_color=BORDER_STRONG,
        left_indent=0.08,
        right_indent=0.08,
    )

    add_heading(doc, "7. Aceite", 1)
    add_para(
        doc,
        "Ao assinar, o cliente confirma o interesse nas condições comerciais acima. O escopo "
        "técnico detalhado das integrações será anexado à ordem de serviço antes do desenvolvimento.",
        size=10.5, color=INK_MUTED, after=14,
    )

    acceptance = doc.add_table(rows=3, cols=2)
    set_table_geometry(acceptance, [4680, 4680])
    labels = [
        ("Pro Distribuidora", "Olume"),
        ("Nome e cargo", "Responsável comercial"),
        ("Data e assinatura", "Data e assinatura"),
    ]
    for ri, (left, right) in enumerate(labels):
        for ci, value in enumerate([left, right]):
            cell = acceptance.cell(ri, ci)
            set_cell_margins(cell, top=130, bottom=130, start=130, end=130)
            set_cell_border(cell, color=BORDER)
            set_cell_text(cell, value, size=9.5,
                          color=PRIMARY_DARK if ri == 0 else INK_MUTED,
                          bold=(ri == 0))
            if ri == 0:
                set_cell_shading(cell, PRIMARY_SOFT)

    add_para(doc, "Conversas que permanecem acesas.", size=13, color=PRIMARY,
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=20, after=3)
    add_para(doc, "Olume Chat  |  olumechat.com.br", size=9.5, color=INK_MUTED,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=0, line_spacing=1.0)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
